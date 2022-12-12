from docx import Document
import pandas
import sqlite3 as sqlite
import urllib.request
from json import loads
from pubobjects import publicacion


def excel_to_db(filename, vconn):
    def update_orcid(filename):
        def get_doi_orcid(orcidurl):
            req = urllib.request.Request(orcidurl)
            req.add_header('Accept', 'application/json')
            with urllib.request.urlopen(req, timeout=15) as f:
                jsonbruto = f.read()
                json = loads(jsonbruto.decode("utf-8"))
                # print(json)
                listadoi = []
                if 'activities-summary' in json.keys():
                    activities = json['activities-summary']
                    for works in activities['works']['group']:
                        for summary in works['work-summary']:
                            if 'external-ids' in summary.keys():
                                for ids in summary['external-ids']:
                                    for id in summary['external-ids']['external-id']:
                                        if (id.get('external-id-type')) != None:
                                            if id.get('external-id-type') == 'doi':
                                                url = id.get('external-id-value')
                                                if "doi.org" in url:
                                                    url = url.split(".org/")[1]
                                                urldoi = 'http://dx.doi.org/' + url
                                                if urldoi not in listadoi:
                                                    listadoi.append(urldoi)
            return listadoi

        academicos = pandas.read_excel(filename, sheet_name='Base_Acad')
        academicos.replace(u'\xa0', u'', regex=True, inplace=True)
        academicos['email'] = academicos['email'].str.lower()
        academicos.dropna(how='all')
        filtered_academicos = academicos[academicos['Orcid'].notnull()]
        df2 = filtered_academicos[['email', 'Orcid']].copy()
        listapub = []
        listamaestro = []
        for index, row in df2.iterrows():
            print('Actualizando doi publicaciones via Orcid de: ' + row["email"])
            listadoi = get_doi_orcid(row["Orcid"])
            print(str(len(listadoi)) + ' publicaciones')
            for doi in listadoi:
                listapub.append([row["email"], doi])

        publicaciones = pandas.DataFrame(listapub, columns=['email', 'doi'])

        with pandas.ExcelWriter(filename, mode='a', if_sheet_exists='replace') as writer:
            publicaciones.to_excel(writer, sheet_name='publicaciones', index=False)

    def to_sql(vfilename, vsheetname, vtable_name, vconn2):
        df = pandas.read_excel(vfilename, sheet_name=vsheetname)
        df.replace(u'\xa0', u'', regex=True, inplace=True)
        df['email'] = df['email'].str.lower()
        df.dropna(how='all')
        df.to_sql(vtable_name, vconn2, if_exists='replace', index=False)

    update_orcid(filename)
    print('Generando base de datos desde excel')
    to_sql(filename, 'Base_Acad', 'base_acad', vconn)
    to_sql(filename, 'tesis_postgrado', 'tesis', vconn)
    to_sql(filename, 'publicaciones', 'publicaciones', vconn)
    to_sql(filename, 'proyectos', 'proyectos', vconn)
    to_sql(filename, 'consultorias', 'consultorias', vconn)
    print('Base de datos actualizada desde excel')
    print('Obteniendo datos de publicaciones')


def add_table_tesis(cell, vconn, id_prof, tipo_tesis):
    def setuptable(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['Año', 'Autor', 'Titulo de tesis', 'Nombre del programa', 'Institución']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    cell.add_paragraph('Como guía tesis')
    table_t1 = cell.add_table(rows=1, cols=5)
    setuptable(table_t1)
    cur = vconn.cursor()
    cur.execute(
        "SELECT b.anno, b.autor, b.titulo, b.programa, b.institucion FROM tesis b where b.tipo =? and b.rol='Guia' and b.email=?",
        [tipo_tesis, id_prof])
    rows = cur.fetchall()
    for row in rows:
        row_cells = table_t1.add_row().cells
        row_cells[0].text = str(row['anno'])
        row_cells[1].text = row['autor']
        row_cells[2].text = row['titulo']
        row_cells[3].text = row['programa']
        row_cells[4].text = row['institucion']

    cell.add_paragraph('Como co-guía tesis')
    table_t2 = cell.add_table(rows=1, cols=5)
    setuptable(table_t2)
    cur.execute(
        "SELECT b.anno, b.autor, b.titulo, b.programa, b.institucion FROM tesis b where b.tipo =? and b.rol='Co-Guia' and b.email=?",
        [tipo_tesis, id_prof])
    rows = cur.fetchall()
    for row in rows:
        row_cells = table_t2.add_row().cells
        row_cells[0].text = str(row['anno'])
        row_cells[1].text = row['autor']
        row_cells[2].text = row['titulo']
        row_cells[3].text = row['programa']
        row_cells[4].text = row['institucion']
    cur.close()


def add_table_proyectos(cell):
    def setuptable(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['Título', 'Fuente de financiamiento', 'Año de adjudicación', 'Periodo de ejecución', 'Rol']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    table_t = cell.add_table(rows=1, cols=5)
    setuptable(table_t)


def add_table_consultorias(cell):
    def setuptable(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['Título', 'Institucion', 'Año de adjudicación', 'Periodo de ejecución', 'Objetivo']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    table_t = cell.add_table(rows=1, cols=5)
    setuptable(table_t)


def add_table_publicaciones(cell, vconn, id_prof):
    def setuptable(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['N°', 'Autores', 'Año', 'Titulo del articulo', 'Nombre revista', 'Estado', 'ISSN',
                    'Factor de impacto']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    def setuptablelibro(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['N°', 'Autores', 'Año', 'Titulo del capitulo/libro', 'Lugar', 'Editorial']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    def setuptableotro(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['N°', 'Autores', 'Año', 'Titulo de publicación', 'Lugar', 'Editorial', 'Estado', 'Observaciones']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    def setuptablepatentes(table_v):
        table_v.style = 'TableGrid'
        cabecera = ['N°', 'Inventores', 'Nombre Patente', 'Fecha solicitud', 'Fecha publicación', 'N° registro',
                    'Estado']
        for c in range(len(cabecera)):
            table_v.cell(0, c).text = cabecera[c]

    cell.add_paragraph('Publicaciones indexadas')
    table_t_pub = cell.add_table(rows=1, cols=8)
    setuptable(table_t_pub)
    cur = vconn.cursor()
    cur.execute(
        "select p.email,m.autores,m.anno,m.titulo,ifnull(m.revista,'Revista')||'('||ifnull(m.ref_revista,'')||') ' ||m.doi as revista,m.isbn,m.factor from publicaciones p, master_doi m where p.doi=m.doi and p.email=?",
        [id_prof])
    rows = cur.fetchall()
    i = 1
    for row in rows:
        row_cells = table_t_pub.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = row['autores']
        row_cells[2].text = str(row['anno'])
        row_cells[3].text = row['titulo']
        row_cells[4].text = row['revista']
        row_cells[5].text = 'Publicada'
        row_cells[6].text = row['isbn']
        row_cells[7].text = row['factor']
        i = i + 1
    cur.close()

    cell.add_paragraph('Libros y capítulos de libro ')
    table_t = cell.add_table(rows=1, cols=6)
    setuptablelibro(table_t)
    cell.add_paragraph('Otras publicaciones ')
    table_t = cell.add_table(rows=1, cols=8)
    setuptableotro(table_t)
    cell.add_paragraph('Patentes ')
    table_t = cell.add_table(rows=1, cols=7)
    setuptablepatentes(table_t)


def get_publicaciones(vconn, id_prof):
    def get_json(url, vconn):
        vurl = ''.join(url.splitlines())
        cursor = vconn.execute('select p.json from master_doi p where doi=? and p.json is not null', [vurl])
        for fila in cursor:
            return fila[0]

    def push_json(url, vjson, vconn):
        vurl = ''.join(url.splitlines())
        json_int = loads(vjson.decode("utf-8"))
        try:
            pub = publicacion(json_int['title'], url)
            pub.add_authors(json_int['author'])
            pub.add_volumen(json_int)
            pub.issn = json_int["ISSN"]
            pub.anno = str(json_int['published']['date-parts'][0][0])
            pub.journal = json_int["container-title"]
            issn, impact, Q = journal_issn_search(pub.issn, conn)
            pub.issn = str(issn)
            pub.impact = f'{impact:.3f} ({Q})'
            pub.found = True
            print(pub.title)
        except Exception as e:
            print(str(e) + " - Error!")
            raise (e)
        lista_autores=pub.get_autorlist(True)
        vconn.execute('insert into master_doi (json,doi,autores, anno, titulo, revista, ref_revista, isbn, factor) values (?,?,?,?,?,?,?,?,?)',
                      [vjson, vurl,lista_autores,pub.anno,pub.title,pub.journal,pub.vol,pub.issn,pub.impact])
        vconn.commit()
    def journal_issn_search(journalissn, vconn):
        if len(journalissn) == 1:
            cursor = vconn.execute(
                'select w.ISSN,ifnull(w.IF_2022,0) if_2022,w.JIF_Quartile from WOS w where w.EISSN=? or w.ISSN=? or w.EISSN=? or w.ISSN=?',
                (json["ISSN"][0], json["ISSN"][0], json["ISSN"][0], json["ISSN"][0]))
        else:
            cursor = vconn.execute(
                'select w.ISSN,ifnull(w.IF_2022,0) if_2022,w.JIF_Quartile from WOS w where w.EISSN=? or w.ISSN=? or w.EISSN=? or w.ISSN=?',
                (json["ISSN"][1], json["ISSN"][1], json["ISSN"][0], json["ISSN"][0]))
        for fila in cursor:
            return fila[0], fila[1], fila[2]
        print("ISSN Not in DB")
        return journalissn[0], 0.0, 'n/a'

    cur = vconn.cursor()
    cur.execute("select p.email,p.doi from publicaciones p where p.email=?", [id_prof])
    rows = cur.fetchall()
    for row in rows:
        url = row['doi']
        try:
            print("Connecting!")
            json = None
            arreglo = get_json(url, vconn)
            if arreglo is not None:
                json = loads(arreglo.decode("utf-8"))
                print("Response from db")
            else:
                try:
                    req = urllib.request.Request(url)
                    req.add_header('Accept', 'application/json')
                    with urllib.request.urlopen(req, timeout=15) as f:
                        jsonbruto = f.read()
                        json = loads(jsonbruto.decode("utf-8"))
                    print("Response from web")
                    push_json(url, jsonbruto, vconn)
                    print("Set into db")
                except Exception as e:
                    print('Error in web ' + str(e))
                    raise (e)

        except Exception as e:
            print(str(e))


def db_2_doc(filename, vconn):
    print('Base actualizada')
    vconn.row_factory = sqlite.Row
    cur = vconn.cursor()
    cur.execute(
        "SELECT b.email,ifnull(b.linea_invest,'') as linea_invest,b.rut,b.nombre,b.tipo,b.profesion||';'||ifnull(b.inst_profesion,'Universidad')||';'||ifnull(b.pais_profesion,'Sin Info') as profesion, b.max_grado||';'||ifnull(b.institucion_grado,'')||';'||ifnull(round(b.ano_max_grado,0),'')||';'||ifnull(b.pais_grado,'') as grado FROM base_acad b where b.tipo ='Nucleo'")
    rows = cur.fetchall()
    for row in rows:
        print('Escribiendo archivo ' + row['nombre'])
        document = Document()
        table = document.add_table(rows=12, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        table.cell(0, 0).text = 'Nombre del académico'
        table.cell(0, 1).text = row['nombre']
        table.cell(1, 0).text = 'Carácter del vínculo (claustro/núcleo, colaborador o visitante'
        table.cell(1, 1).text = row['tipo']
        table.cell(2, 0).text = 'Título profesional,  institución, país'
        table.cell(2, 1).text = row['profesion']
        table.cell(3, 0).text = 'Grado académico máximo (especificar área disciplinar), institución, año de graduación y país '
        table.cell(3, 1).text = row['grado']
        table.cell(4, 0).text = 'Línea(s) de investigación'
        table.cell(4, 1).text = row['linea_invest']

        table.cell(5, 0).text = 'Tesis de magíster  dirigidas en los últimos 10 años (finalizadas)'
        add_table_tesis(table.cell(5, 1), vconn, row['email'], 'Magister')

        table.cell(6, 0).text = 'Tesis de doctorado dirigidas en los últimos 10 años (finalizadas)'
        add_table_tesis(table.cell(6, 1), vconn, row['email'], 'Doctorado')

        table.cell(7, 0).text = 'Listado de publicaciones'
        get_publicaciones(vconn, row['email'])
        add_table_publicaciones(table.cell(7, 1), vconn, row['email'])

        table.cell(8, 0).text = 'Listado de proyectos de investigación  en los últimos 10 años'
        add_table_proyectos(table.cell(8, 1))

        table.cell(9, 0).text = 'Listado de proyectos de intervención, innovación y/o desarrollo tecnológico'
        add_table_proyectos(table.cell(9, 1))

        table.cell(10,
                   0).text = 'Listado de consultorías y/o  asistencias técnicas, en calidad de responsable, en los últimos 10 años'
        add_table_consultorias(table.cell(10, 1))

        document.save('./output/' + row['rut'] + '_' + row['nombre'] + '.docx')


base_file = 'Base_Academicos_demo.xlsx'
db_academics = './bd_academic.sqlite'
conn = sqlite.connect(db_academics)
excel_to_db(base_file, conn)
db_2_doc(base_file, conn)
